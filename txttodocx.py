from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import unicodedata
# -*- coding: utf8 -*-

# importing file
output = open("output.txt", "w")
inputFile = open("input.txt", "r", encoding='UTF8')

# checking if test.docx exsists and if so delete it
if os.path.exists('test.docx'):
    os.remove('test.docx')

#creating new docx
document = Document()
document.save('text.docx')

# TODO: checking if it has correct formatting on all the lines
# correct formatting: "word=word^p"

# sorting Document
lines = inputFile.readlines()
lines.sort()
numOfLines = len(lines)

lineCounter = 0;
howMany = 0

print('How many copies (multiples of 3): ')
howMany = input()
howMany = int(howMany)

#creating a table
table = document.add_table(rows = howMany, cols = 3)
table.style = 'Table Grid'

for x in range(0, howMany):
	for y in range(3):
		cell = table.cell(x,y)
		paragraph = cell.paragraphs
		p = paragraph[0]
		for line in lines:
			lineCounter = lineCounter + 1
			line = line.strip('\n')
			line = line.strip('\t')
			split = line.split('=')
			p.add_run(split[0]).bold = True
			p.add_run('=')
			p.add_run(split[1])
			p.add_run('; ')
			if int(numOfLines/2) == lineCounter:
				p.add_run('\n\n')
		#formatting paragraph inside the table
		format = p.paragraph_format
		style = document.styles['Normal']
		font = style.font
		format.line_spacing = 0.8
		font.name = 'Arial'
		font.size = Pt(3)
		p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
		cell.bottom_margin = 0.2
		cell.left_padding = 0.05
		cell.right_padding = 0.05
		lineCounter = 0

# saving the document
document.save('text.docx')

inputFile.close()
output.close()